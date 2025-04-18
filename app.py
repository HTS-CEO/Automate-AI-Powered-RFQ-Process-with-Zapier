from flask import Flask, render_template_string, request, jsonify
import os
import PyPDF2
from docx import Document
import openai
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import tempfile
from werkzeug.utils import secure_filename

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
openai.api_key = 'your-openai-api-key'
SMTP_SERVER = 'your-smtp-server.com'
SMTP_PORT = 587
SMTP_USERNAME = 'your-email@company.com'
SMTP_PASSWORD = 'your-email-password'

HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <title>RFQ Automation Tool</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }
        .container {
            background-color: white;
            padding: 30px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        h1 {
            color: #2c3e50;
            text-align: center;
        }
        .upload-section {
            border: 2px dashed #3498db;
            padding: 20px;
            text-align: center;
            margin-bottom: 20px;
            border-radius: 5px;
        }
        .btn {
            background-color: #3498db;
            color: white;
            padding: 10px 15px;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            font-size: 16px;
        }
        .btn:hover {
            background-color: #2980b9;
        }
        #file-info {
            margin-top: 10px;
            font-size: 14px;
            color: #7f8c8d;
        }
        #results {
            margin-top: 20px;
            padding: 15px;
            background-color: #f8f9fa;
            border-radius: 5px;
            display: none;
        }
        .status {
            margin-top: 20px;
            padding: 10px;
            border-radius: 4px;
        }
        .success {
            background-color: #d4edda;
            color: #155724;
        }
        .error {
            background-color: #f8d7da;
            color: #721c24;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>RFQ Automation Tool</h1>
        
        <div class="upload-section">
            <h3>Upload RFQ Source Document</h3>
            <input type="file" id="fileInput" accept=".pdf,.docx">
            <p id="file-info">No file selected</p>
            <button class="btn" onclick="processDocument()">Process Document</button>
        </div>
        
        <div id="status" class="status" style="display:none;"></div>
        
        <div id="results">
            <h3>Extracted Information:</h3>
            <pre id="extracted-data"></pre>
            <button class="btn" onclick="generateRFQ()">Generate RFQ</button>
        </div>
    </div>

    <script>
        document.getElementById('fileInput').addEventListener('change', function(e) {
            const fileInfo = document.getElementById('file-info');
            if (this.files.length > 0) {
                fileInfo.textContent = `Selected file: ${this.files[0].name}`;
            } else {
                fileInfo.textContent = 'No file selected';
            }
        });

        async function processDocument() {
            const fileInput = document.getElementById('fileInput');
            const statusDiv = document.getElementById('status');
            
            if (fileInput.files.length === 0) {
                statusDiv.textContent = 'Please select a file first';
                statusDiv.className = 'status error';
                statusDiv.style.display = 'block';
                return;
            }
            
            statusDiv.textContent = 'Processing document...';
            statusDiv.className = 'status';
            statusDiv.style.display = 'block';
            
            const formData = new FormData();
            formData.append('file', fileInput.files[0]);
            
            try {
                const response = await fetch('/process', {
                    method: 'POST',
                    body: formData
                });
                
                const result = await response.json();
                
                if (response.ok) {
                    statusDiv.textContent = 'Document processed successfully!';
                    statusDiv.className = 'status success';
                    document.getElementById('extracted-data').textContent = JSON.stringify(result.extracted_data, null, 2);
                    document.getElementById('results').style.display = 'block';
                } else {
                    statusDiv.textContent = 'Error: ' + result.error;
                    statusDiv.className = 'status error';
                }
            } catch (error) {
                statusDiv.textContent = 'Error: ' + error.message;
                statusDiv.className = 'status error';
            }
        }
        
        async function generateRFQ() {
            const statusDiv = document.getElementById('status');
            statusDiv.textContent = 'Generating RFQ...';
            statusDiv.className = 'status';
            statusDiv.style.display = 'block';
            
            try {
                const response = await fetch('/generate-rfq', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json',
                    },
                    body: JSON.stringify({
                        extracted_data: document.getElementById('extracted-data').textContent
                    })
                });
                
                if (response.ok) {
                    statusDiv.textContent = 'RFQ generated and sent for review!';
                    statusDiv.className = 'status success';
                } else {
                    const result = await response.json();
                    statusDiv.textContent = 'Error: ' + result.error;
                    statusDiv.className = 'status error';
                }
            } catch (error) {
                statusDiv.textContent = 'Error: ' + error.message;
                statusDiv.className = 'status error';
            }
        }
    </script>
</body>
</html>
"""

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(filepath):
    text = ""
    with open(filepath, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(filepath):
    doc = Document(filepath)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_rfq_data(text):
    prompt = f"""Extract the following information from this RFQ document:
    - Product/service descriptions
    - Quantities
    - Timelines
    - Special requirements
    - Contact information
    
    Return the data in JSON format.
    
    Document text:
    {text}"""
    
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that extracts RFQ information from documents."},
            {"role": "user", "content": prompt}
        ]
    )
    
    return response.choices[0].message.content

def generate_rfq_document(data):
    # This would use your company's RFQ template
    # For simplicity, we'll just create a basic text document
    rfq_content = f"""REQUEST FOR QUOTE
    
Vendor Information:
{data.get('vendor_info', 'N/A')}

Products/Services:
{data.get('products', 'N/A')}

Quantities:
{data.get('quantities', 'N/A')}

Timeline:
{data.get('timeline', 'N/A')}

Special Requirements:
{data.get('requirements', 'N/A')}
"""
    
    return rfq_content

def send_email(to_email, subject, body, attachment_content=None, attachment_filename=None):
    msg = MIMEMultipart()
    msg['From'] = SMTP_USERNAME
    msg['To'] = to_email
    msg['Subject'] = subject
    
    msg.attach(MIMEText(body, 'plain'))
    
    if attachment_content and attachment_filename:
        part = MIMEApplication(attachment_content, Name=attachment_filename)
        part['Content-Disposition'] = f'attachment; filename="{attachment_filename}"'
        msg.attach(part)
    
    with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
        server.starttls()
        server.login(SMTP_USERNAME, SMTP_PASSWORD)
        server.send_message(msg)

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/process', methods=['POST'])
def process_document():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        file.save(filepath)
        
        try:
            if filename.endswith('.pdf'):
                text = extract_text_from_pdf(filepath)
            elif filename.endswith('.docx'):
                text = extract_text_from_docx(filepath)
            
            extracted_data = extract_rfq_data(text)
            return jsonify({
                'status': 'success',
                'extracted_data': extracted_data
            })
        except Exception as e:
            return jsonify({'error': str(e)}), 500
        finally:
            if os.path.exists(filepath):
                os.remove(filepath)
    else:
        return jsonify({'error': 'Invalid file type'}), 400

@app.route('/generate-rfq', methods=['POST'])
def generate_rfq():
    try:
        data = request.json.get('extracted_data')
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        # Parse the extracted data (in a real app, you'd have proper validation)
        rfq_data = eval(data) if isinstance(data, str) else data
        
        # Generate the RFQ document
        rfq_content = generate_rfq_document(rfq_data)
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.txt') as tmp:
            tmp.write(rfq_content.encode('utf-8'))
            tmp_path = tmp.name
        
        # Send email with attachment
        send_email(
            to_email='procurement@yourcompany.com',
            subject='New RFQ for Review',
            body='Please review the attached RFQ document.',
            attachment_content=rfq_content,
            attachment_filename='rfq_draft.txt'
        )
        
        # Clean up
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
